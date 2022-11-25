from datetime import datetime
from docx import Document
from docx.shared import Pt


def get_current_date():
    """
    取得當前的日期，並整理成文件命名與文件標題所需的格式
    """
    current_datetime = datetime.now()
    docs_suffix = f"{current_datetime.year}{current_datetime.month}{current_datetime.day}"
    docs_title_time = f"{current_datetime.month}/{current_datetime.day}"
    return docs_suffix, docs_title_time


if __name__ == "__main__":
    _docs_suffix, _docs_title_time = get_current_date()
    folder = r"C:\Users\jasper chiu\OneDrive - 百通科技股份有限公司"
    docs_path = folder + rf"\Jasper_工作日誌_{_docs_suffix}.docx"

    # if "檢測該資料是否存在，若存在則不往下處理"
    # 建立物件(word)
    document = Document()
    # 添加段落文字
    p1 = document.add_paragraph()
    text1 = p1.add_run(f"工作日誌 {_docs_title_time}")  # 首行段落內容
    text1.font.size = Pt(20)  # 設定字體大小
    p2 = document.add_paragraph()  # 次行段落內容
    text2 = p2.add_run("\n")  # 空一行方便後續操作
    text2.font.size = Pt(18)
    document.save(docs_path)
